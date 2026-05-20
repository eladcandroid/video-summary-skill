#!/usr/bin/env python3
"""write_docx.py --input <md> --title <str> --output <docx>

Convert a Hebrew markdown-ish lesson file to a right-to-left .docx.

Syntax recognised in the input:
- `# Heading 1` (one line)
- `## Heading 2`
- `### Heading 3`
- Blank line = paragraph break.
- Everything else = body paragraph (joined within a block by single \\n,
  separated between blocks by blank lines).
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document  # factory function, not a class
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

HEBREW_FONT = "David"
FALLBACK_FONT = "Arial"
BODY_SIZE_PT = 12
HEADING_SIZES = {1: 20, 2: 16, 3: 13}


def _set_rtl(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _set_run_rtl(run, size_pt: int) -> None:
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rPr.append(rtl)

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FALLBACK_FONT)
    rFonts.set(qn("w:hAnsi"), FALLBACK_FONT)
    rFonts.set(qn("w:cs"), HEBREW_FONT)
    rPr.append(rFonts)

    sz = OxmlElement("w:szCs")
    sz.set(qn("w:val"), str(size_pt * 2))
    rPr.append(sz)

    run.font.size = Pt(size_pt)


def _set_doc_language(doc) -> None:
    normal = doc.styles["Normal"]
    rpr = normal.element.get_or_add_rPr()
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "he-IL")
    lang.set(qn("w:bidi"), "he-IL")
    rpr.append(lang)
    normal.font.name = FALLBACK_FONT
    normal.font.size = Pt(BODY_SIZE_PT)


def _add_heading(doc, text: str, level: int) -> None:
    p = doc.add_heading(level=level)
    _set_rtl(p)
    run = p.add_run(text)
    _set_run_rtl(run, HEADING_SIZES.get(level, BODY_SIZE_PT))
    run.bold = True


def _add_paragraph(doc, text: str) -> None:
    p = doc.add_paragraph()
    _set_rtl(p)
    run = p.add_run(text)
    _set_run_rtl(run, BODY_SIZE_PT)


def _parse_blocks(text: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    buf: list[str] = []

    def flush_paragraph():
        if buf:
            blocks.append(("p", "\n".join(buf).strip()))
            buf.clear()

    for raw in text.splitlines():
        line = raw.rstrip()
        if not line:
            flush_paragraph()
            continue
        if line.startswith("### "):
            flush_paragraph()
            blocks.append(("h3", line[4:].strip()))
        elif line.startswith("## "):
            flush_paragraph()
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("# "):
            flush_paragraph()
            blocks.append(("h1", line[2:].strip()))
        else:
            buf.append(line)
    flush_paragraph()
    return blocks


def build_document(input_text: str, fallback_title: str):
    doc = Document()
    _set_doc_language(doc)

    section = doc.sections[0]
    sectPr = section._sectPr
    bidi = OxmlElement("w:bidi")
    sectPr.append(bidi)

    blocks = _parse_blocks(input_text)

    if not blocks or blocks[0][0] != "h1":
        _add_heading(doc, fallback_title, 1)

    for kind, payload in blocks:
        if kind == "h1":
            _add_heading(doc, payload, 1)
        elif kind == "h2":
            _add_heading(doc, payload, 2)
        elif kind == "h3":
            _add_heading(doc, payload, 3)
        else:
            _add_paragraph(doc, payload)

    return doc


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", required=True, help="Markdown-ish edited lesson")
    ap.add_argument("--title", required=True, help="Fallback title if input has no H1")
    ap.add_argument("--output", required=True, help="Output .docx path")
    args = ap.parse_args()

    in_path = Path(args.input)
    out_path = Path(args.output)

    if not in_path.exists():
        print(f"input not found: {in_path}", file=sys.stderr)
        return 1

    text = in_path.read_text(encoding="utf-8")
    doc = build_document(text, args.title)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(str(out_path.resolve()))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
